"""
Enhanced Document Processing Module for RAG System
Supports: PDF, TXT, DOCX, XLSX text extraction, tables, OCR, and image understanding
"""
import sys
import os
import io
import base64
from typing import List, Dict, Optional, Tuple
from pathlib import Path
# PDF Processing
import pdfplumber
import fitz  # PyMuPDF for fallback
from pdf2image import convert_from_path
import camelot  # For complex table extraction
# Office Document Processing
import docx
import pandas as pd

# OCR
from PIL import Image
import pytesseract

# Image Understanding
import torch
from transformers import Blip2Processor, Blip2ForConditionalGeneration

# LangChain
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .config import RAGConfig


def _safe_print(*args, **kwargs):
    """Print with Unicode-safe fallback for Windows consoles."""
    try:
        print(*args, **kwargs)
    except UnicodeEncodeError:
        text = ' '.join(str(a) for a in args)
        print(text.encode(sys.stdout.encoding or 'utf-8', errors='replace').decode(
            sys.stdout.encoding or 'utf-8', errors='replace'), **kwargs)


class EnhancedDocumentProcessor:
    """
    Advanced document processor with multi-modal capabilities:
    - Text extraction from PDF, TXT, DOCX, and XLSX
    - Table extraction with Camelot/pdfplumber/pandas
    - OCR for scanned PDF pages with Tesseract
    - Image understanding with BLIP-2
    """
    
    def __init__(self, config: RAGConfig = None):
        """
        Initialize enhanced document processor
        
        Args:
            config: RAGConfig instance
        """
        self.config = config or RAGConfig()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.CHUNK_SIZE,
            chunk_overlap=self.config.CHUNK_OVERLAP,
            length_function=len,
            separators=self.config.TEXT_SEPARATORS
        )
        
        # Image understanding model (lazy loaded)
        self.blip_processor = None
        self.blip_model = None
        self.device = 'cuda' if torch.cuda.is_available() else 'cpu'
        
        # Processing statistics
        self.stats = {
            'total_pages': 0,
            'text_pages': 0,
            'ocr_pages': 0,
            'tables_extracted': 0,
            'images_processed': 0
        }
    
    def load_image_model(self):
        """Lazy load BLIP-2 model for image understanding"""
        if self.blip_model is None:
            _safe_print("🔄 Loading BLIP-2 model for image understanding...")
            
            model_name = "Salesforce/blip2-opt-2.7b"  # Smaller model for efficiency
            
            self.blip_processor = Blip2Processor.from_pretrained(model_name)
            self.blip_model = Blip2ForConditionalGeneration.from_pretrained(
                model_name,
                torch_dtype=torch.float16 if self.device == 'cuda' else torch.float32,
                device_map="auto"
            )
            
            _safe_print(f"✅ BLIP-2 model loaded on {self.device}")
    
    def is_page_scanned(self, page) -> bool:
        """Detect if a page is scanned/image-based"""
        text = page.extract_text()
        
        # If very little text but has images, likely scanned
        if not text or len(text.strip()) < 50:
            if len(page.images) > 0:
                return True
        
        # Check text-to-image ratio
        if text and len(page.images) > 0:
            text_len = len(text.strip())
            page_area = page.width * page.height
            
            # If very little text relative to page size, might be scanned
            if text_len < 100 and page_area > 100000:
                return True
        
        return False
    
    def extract_text_with_pdfplumber(self, pdf_path: str) -> List[Dict]:
        """Extract text from PDF using pdfplumber"""
        pages_data = []
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Try regular text extraction
                text = page.extract_text()
                
                # Check if page needs OCR
                needs_ocr = self.is_page_scanned(page)
                
                if needs_ocr or not text or len(text.strip()) < 20:
                    _safe_print(f"   Page {page_num}: Applying OCR (scanned/low text)")
                    text = self.ocr_page(pdf_path, page_num)
                    self.stats['ocr_pages'] += 1
                else:
                    self.stats['text_pages'] += 1
                
                pages_data.append({
                    'page_number': page_num,
                    'text': text.strip() if text else "",
                    'needs_ocr': needs_ocr,
                    'char_count': len(text) if text else 0,
                    'word_count': len(text.split()) if text else 0,
                    'has_images': len(page.images) > 0,
                    'image_count': len(page.images)
                })
                
                self.stats['total_pages'] += 1
        
        return pages_data
    
    def ocr_page(self, pdf_path: str, page_num: int) -> str:
        """Perform OCR on a specific page using Tesseract"""
        try:
            images = convert_from_path(
                pdf_path,
                first_page=page_num + 1,
                last_page=page_num + 1,
                dpi=300
            )
            
            if not images:
                return ""
            
            text = pytesseract.image_to_string(
                images[0],
                lang='eng',
                config='--psm 1'
            )
            
            return text.strip()
            
        except Exception as e:
            _safe_print(f"      OCR failed for page {page_num}: {e}")
            return ""
    
    def extract_tables_camelot(self, pdf_path: str, page_num: int) -> List[str]:
        """Extract tables using Camelot"""
        try:
            tables = camelot.read_pdf(
                pdf_path,
                pages=str(page_num + 1),
                flavor='lattice',
                suppress_stdout=True
            )
            
            table_texts = []
            for table in tables:
                df = table.df
                markdown_table = df.to_markdown(index=False)
                table_texts.append(markdown_table)
                self.stats['tables_extracted'] += 1
            
            return table_texts
            
        except Exception as e:
            return []
    
    def extract_tables_pdfplumber(self, page) -> List[str]:
        """Extract tables using pdfplumber (fallback)"""
        try:
            tables = page.extract_tables()
            
            if not tables:
                return []
            
            table_texts = []
            for table in tables:
                if not table or len(table) == 0:
                    continue
                
                table_str = "\n".join([" | ".join([str(cell) if cell else "" for cell in row]) for row in table])
                table_texts.append(f"\n[TABLE]\n{table_str}\n[/TABLE]\n")
                self.stats['tables_extracted'] += 1
            
            return table_texts
            
        except Exception as e:
            return []
    
    def extract_images_and_describe(self, pdf_path: str, page_num: int) -> List[Dict]:
        """Extract images from page and generate descriptions using BLIP-2"""
        image_descriptions = []
        
        try:
            images = convert_from_path(
                pdf_path,
                first_page=page_num + 1,
                last_page=page_num + 1,
                dpi=150
            )
            
            if not images:
                return []
            
            page_image = images[0]
            self.load_image_model()
            
            # Caption
            inputs = self.blip_processor(page_image, return_tensors="pt").to(self.device)
            generated_ids = self.blip_model.generate(**inputs, max_new_tokens=50)
            caption = self.blip_processor.batch_decode(generated_ids, skip_special_tokens=True)[0].strip()
            
            # Detailed description
            prompt = "Question: What is shown in this image? Describe the key elements. Answer:"
            inputs = self.blip_processor(page_image, text=prompt, return_tensors="pt").to(self.device)
            generated_ids = self.blip_model.generate(**inputs, max_new_tokens=100)
            description = self.blip_processor.batch_decode(generated_ids, skip_special_tokens=True)[0].strip()
            
            image_descriptions.append({
                'caption': caption,
                'description': description,
                'page': page_num
            })
            
            self.stats['images_processed'] += 1
            
        except Exception as e:
            _safe_print(f"      Image processing failed for page {page_num}: {e}")
        
        return image_descriptions
    
    def process_pdf_enhanced(self, pdf_path: str, source_name: str,
                             extract_tables: bool = True,
                             describe_images: bool = True) -> List[Dict]:
        """Enhanced PDF processing with all features"""
        enhanced_pages = []
        pages_data = self.extract_text_with_pdfplumber(pdf_path)
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_idx, page_data in enumerate(pages_data):
                page = pdf.pages[page_idx]
                page_content = page_data['text']
                
                if extract_tables:
                    tables = self.extract_tables_camelot(pdf_path, page_idx)
                    if not tables:
                        tables = self.extract_tables_pdfplumber(page)
                    
                    if tables:
                        # Normalize formatting to make sure smart-chunking catches it
                        formatted_tables = [t if "[TABLE]" in t else f"\n[TABLE]\n{t}\n[/TABLE]\n" for t in tables]
                        page_content += "\n\n" + "\n\n".join(formatted_tables)
                
                if describe_images and page_data['has_images']:
                    image_descriptions = self.extract_images_and_describe(pdf_path, page_idx)
                    for img_desc in image_descriptions:
                        page_content += f"\n\n[IMAGE DESCRIPTION: {img_desc['description']}]"
                
                enhanced_pages.append({
                    'page_number': page_idx,
                    'source': source_name,
                    'text': page_content,
                    'char_count': len(page_content),
                    'word_count': len(page_content.split()),
                    'needs_ocr': page_data['needs_ocr'],
                    'has_tables': extract_tables and len(self.extract_tables_pdfplumber(page)) > 0,
                    'has_images': page_data['has_images']
                })
        
        return enhanced_pages

    def process_txt(self, file_path: str, source_name: str) -> List[Dict]:
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
                
        return [{
            'page_number': 0,
            'source': source_name,
            'text': text,
            'char_count': len(text),
            'word_count': len(text.split()),
            'needs_ocr': False,
            'has_tables': False,
            'has_images': False
        }]

    def process_docx(self, file_path: str, source_name: str, extract_tables: bool = True) -> List[Dict]:
        """Process Word documents, extracting paragraphs and tables."""
        doc = docx.Document(file_path)
        content = []
        has_tables = len(doc.tables) > 0

        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)

        if extract_tables and has_tables:
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    data.append(row_data)
                
                if len(data) > 1:
                    try:
                        df = pd.DataFrame(data[1:], columns=data[0])
                        markdown_table = df.to_markdown(index=False)
                        content.append(f"\n[TABLE]\n{markdown_table}\n[/TABLE]\n")
                        self.stats['tables_extracted'] += 1
                    except Exception as e:
                        _safe_print(f"      Table extraction failed in DOCX: {e}")

        full_text = "\n\n".join(content)

        return [{
            'page_number': 0,
            'source': source_name,
            'text': full_text,
            'char_count': len(full_text),
            'word_count': len(full_text.split()),
            'needs_ocr': False,
            'has_tables': extract_tables and has_tables,
            'has_images': False
        }]

    def process_xlsx(self, file_path: str, source_name: str) -> List[Dict]:
        """Process Excel files by converting sheets directly to Markdown tables."""
        content = []
        
        try:
            excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            for sheet_name, df in excel_data.items():
                content.append(f"### Sheet: {sheet_name}")
                
                df.dropna(how='all', inplace=True)
                df.dropna(axis=1, how='all', inplace=True)
                
                if not df.empty:
                    markdown_table = df.to_markdown(index=False)
                    content.append(f"\n[TABLE]\n{markdown_table}\n[/TABLE]\n")
                    self.stats['tables_extracted'] += 1
                    
        except Exception as e:
            _safe_print(f"      Excel extraction failed: {e}")

        full_text = "\n\n".join(content)

        return [{
            'page_number': 0,
            'source': source_name,
            'text': full_text,
            'char_count': len(full_text),
            'word_count': len(full_text.split()),
            'needs_ocr': False,
            'has_tables': True,
            'has_images': False
        }]

    def convert_to_langchain_documents(self, enhanced_pages: List[Dict]) -> List[Document]:
        """Convert enhanced page data to LangChain Document objects"""
        documents = []
        
        for page_data in enhanced_pages:
            doc = Document(
                page_content=page_data['text'],
                metadata={
                    'source': page_data['source'],
                    'page': page_data['page_number'],
                    'page_number': page_data['page_number'],
                    'needs_ocr': page_data.get('needs_ocr', False),
                    'has_tables': page_data.get('has_tables', False),
                    'has_images': page_data.get('has_images', False)
                }
            )
            documents.append(doc)
        
        return documents
    
    def split_documents_smart(self, documents: List[Document]) -> List[Document]:
            """Smart document splitting that preserves tables and important structures"""
            chunks = []
            
            # 1. ADDED: A global counter that persists across ALL pages/documents
            global_chunk_idx = 0  
            
            for doc in documents:
                text = doc.page_content
                has_table = '[TABLE]' in text
                
                if has_table:
                    parts = text.split('[TABLE]')
                    
                    # Removed the 'enumerate(parts)' here since we use the global counter
                    for part in parts:
                        if '[/TABLE]' in part:
                            table_content = '[TABLE]' + part
                            chunks.append(Document(
                                page_content=table_content,
                                metadata={**doc.metadata, 'chunk_type': 'table', 'chunk_index': global_chunk_idx}
                            ))
                            # 2. ADDED: Increment after appending a table
                            global_chunk_idx += 1
                        else:
                            text_chunks = self.text_splitter.split_text(part)
                            # Removed the 'enumerate(text_chunks)' here
                            for chunk_text in text_chunks:
                                chunks.append(Document(
                                    page_content=chunk_text,
                                    metadata={**doc.metadata, 'chunk_type': 'text', 'chunk_index': global_chunk_idx}
                                ))
                                # 3. ADDED: Increment after appending standard text around tables
                                global_chunk_idx += 1
                else:
                    text_chunks = self.text_splitter.split_documents([doc])
                    # Removed 'enumerate(text_chunks)' here as well
                    for chunk in text_chunks:
                        chunk.metadata['chunk_type'] = 'text'
                        
                        # 4. CHANGED: Use the global counter instead of the loop 'i' index
                        chunk.metadata['chunk_index'] = global_chunk_idx
                        chunks.append(chunk)
                        
                        # 5. ADDED: Increment for standard page text
                        global_chunk_idx += 1
            
            _safe_print(f"📦 Created {len(chunks)} smart chunks from {len(documents)} pages")
            
            return chunks
    
    def process_document_complete(self, file_path: str, 
                                  extract_tables: bool = True,
                                  describe_images: bool = False) -> List[Document]:
        """
        Complete document processing pipeline supporting PDF, TXT, DOCX, XLSX
        """
        path_obj = Path(file_path)
        source_name = path_obj.name
        ext = path_obj.suffix.lower()
        
        # Reset stats
        self.stats = {k: 0 for k in self.stats}
        
        _safe_print(f"\n📄 Processing: {source_name}")
        
        # Step 1: Enhanced extraction routed by file type
        if ext == '.pdf':
            enhanced_pages = self.process_pdf_enhanced(
                file_path, source_name, extract_tables, describe_images
            )
        elif ext == '.txt':
            enhanced_pages = self.process_txt(file_path, source_name)
            self.stats['text_pages'] += 1
            self.stats['total_pages'] += 1
        elif ext in ['.docx', '.doc']:
            enhanced_pages = self.process_docx(file_path, source_name, extract_tables)
            self.stats['text_pages'] += 1
            self.stats['total_pages'] += 1
        elif ext in ['.xlsx', '.xls', '.csv']:
            enhanced_pages = self.process_xlsx(file_path, source_name)
            self.stats['text_pages'] += 1
            self.stats['total_pages'] += 1
        else:
            raise ValueError(f"Unsupported file extension: {ext}. Supported formats are PDF, TXT, DOCX, XLSX.")
        
        # Print statistics
        _safe_print(f"✅ Processed {self.stats['total_pages']} blocks/pages")
        _safe_print(f"   📝 Text pages: {self.stats['text_pages']}")
        if ext == '.pdf':
            _safe_print(f"   🔍 OCR pages: {self.stats['ocr_pages']}")
            _safe_print(f"   🖼️  Images processed: {self.stats['images_processed']}")
        _safe_print(f"   📊 Tables extracted: {self.stats['tables_extracted']}")

        # Step 2: Convert to LangChain documents
        documents = self.convert_to_langchain_documents(enhanced_pages)
        
        # Step 3: Smart chunking
        chunks = self.split_documents_smart(documents)
        
        return chunks
    
    def get_processing_stats(self) -> Dict:
        """Get processing statistics"""
        return self.stats.copy()